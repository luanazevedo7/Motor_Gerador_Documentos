# -*- coding: utf-8 -*-

"""
Motor de Geração de Documentos v1.0
===================================

Este script é um motor flexível para gerar documentos personalizados (.docx) a
partir de uma base de dados em Excel. Ele é projetado para ser facilmente
adaptável a diferentes necessidades, como a criação de termos, contratos,
certificados, convites e mais, através de um painel de configuração simples.

Autor: Luan Azevedo
Data: 21 de Julho de 2025
Versão: 1.0
"""

import pandas as pd
import docx
import os
from datetime import datetime

# ==============================================================================
# --- PAINEL DE CONTROLE (CONFIGURE SEU PROJETO AQUI) ---
# ==============================================================================

# 1. NOMES DE ARQUIVOS E PASTAS
PASTA_DADOS = 'dados'
NOME_ARQUIVO_EXCEL = 'banco_de_dados_artistas.xlsx'
NOME_ABA_EXCEL = 'termo'

PASTA_MODELO = 'modelo'
NOME_ARQUIVO_MODELO = 'modelo_termo.docx'

PASTA_SAIDA = 'termos_gerados'
PREFIXO_ARQUIVO_SAIDA = 'Termo_Compromisso'

# 2. MAPEAMENTO DE DADOS (Excel -> Documento Word)
# Colunas que serão lidas diretamente da planilha e seus placeholders no Word.
MAPEAMENTO_COLUNAS = {
    'Nome do Artista': '[nome_completo]',
    'CPF': '[cpf]',
    'RG': '[rg]',
    'Endereço Completo': '[endereco]',
}

# 3. CONFIGURAÇÃO DA INTERAÇÃO COM O USUÁRIO
# Qual coluna do Excel deve ser usada para identificar os itens na lista?
COLUNA_IDENTIFICADORA = 'Nome do Artista'

# Quais campos o programa deve perguntar ao usuário em tempo de execução?
# O nome do campo será usado como placeholder no Word (ex: [categoria], [valor])
CAMPOS_INTERATIVOS = ['categoria', 'valor']


# ==============================================================================
# --- LÓGICA DO PROGRAMA (GERALMENTE NÃO PRECISA ALTERAR) ---
# ==============================================================================

def obter_data_formatada():
    """Retorna a data atual formatada por extenso em português."""
    now = datetime.now()
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
             "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    return f"{now.day} de {meses[now.month - 1]} de {now.year}"


def motor_gerador():
    """Função principal que orquestra todo o processo de geração de documentos."""
    print(f"--- Iniciando Motor de Geração de Documentos: {PREFIXO_ARQUIVO_SAIDA} ---")

    caminho_excel = os.path.join(PASTA_DADOS, NOME_ARQUIVO_EXCEL)
    caminho_modelo = os.path.join(PASTA_MODELO, NOME_ARQUIVO_MODELO)

    # --- 1. CARREGAR DADOS ---
    try:
        df_dados = pd.read_excel(caminho_excel, sheet_name=NOME_ABA_EXCEL)
        df_dados = df_dados.astype(str)
    except FileNotFoundError:
        print(f"\n[ERRO] Arquivo de dados não encontrado em: '{caminho_excel}'")
        print("Verifique se a estrutura de pastas e os nomes dos arquivos estão corretos.")
        input("\nPressione Enter para sair.")
        return
    except ValueError:
        print(f"\n[ERRO] Aba '{NOME_ABA_EXCEL}' não encontrada na planilha '{NOME_ARQUIVO_EXCEL}'.")
        input("\nPressione Enter para sair.")
        return
    except Exception as e:
        print(f"\n[ERRO] Ocorreu um erro inesperado ao ler o Excel: {e}")
        input("\nPressione Enter para sair.")
        return

    # --- 2. APRESENTAR E SELECIONAR ITENS ---
    try:
        if df_dados.empty:
            print(f"\n[AVISO] A aba '{NOME_ABA_EXCEL}' está vazia. Não há itens para processar.")
            input("\nPressione Enter para sair.")
            return

        print("\n--- Lista de Itens Disponíveis ---")
        for index, row in df_dados.iterrows():
            print(f"{index}: {row[COLUNA_IDENTIFICADORA]}")
    except KeyError as e:
        print(f"\n[ERRO] A coluna identificadora '{e}' não foi encontrada na planilha.")
        input("\nPressione Enter para sair.")
        return

    print("\n--- Seleção de Itens ---")
    input_usuario = input("Digite os números dos itens, separados por vírgula: ")
    try:
        indices_selecionados = [int(i.strip()) for i in input_usuario.split(',')]
    except ValueError:
        print("\n[ERRO] Seleção inválida. Por favor, digite apenas números.")
        input("\nPressione Enter para sair.")
        return

    # --- 3. LOOP DE GERAÇÃO ---
    os.makedirs(PASTA_SAIDA, exist_ok=True)
    total_gerado = 0

    print("\n--- Processando Documentos ---")
    for index in indices_selecionados:
        if index in df_dados.index:
            registro = df_dados.loc[index]
            nome_identificador = registro[COLUNA_IDENTIFICADORA]

            print(f"\n> Preenchendo dados para: {nome_identificador}")

            dados_interativos = {}
            for campo in CAMPOS_INTERATIVOS:
                valor_campo = input(f"  Digite o valor para [{campo}]: ")
                dados_interativos[f'[{campo}]'] = valor_campo

            try:
                documento = docx.Document(caminho_modelo)

                substituicoes = {}
                substituicoes.update({placeholder: registro[col] for col, placeholder in MAPEAMENTO_COLUNAS.items()})
                substituicoes.update(dados_interativos)
                substituicoes['[data_atual]'] = obter_data_formatada()

                # Lógica de substituição aprimorada para preservar formatação
                for placeholder, valor in substituicoes.items():
                    for paragraph in documento.paragraphs:
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(valor))
                    for table in documento.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    if placeholder in paragraph.text:
                                        for run in paragraph.runs:
                                            if placeholder in run.text:
                                                run.text = run.text.replace(placeholder, str(valor))

                nome_arquivo_seguro = ''.join(c for c in nome_identificador if c.isalnum() or c in (' ', '_')).rstrip()
                caminho_saida = os.path.join(PASTA_SAIDA, f"{PREFIXO_ARQUIVO_SAIDA}_{nome_arquivo_seguro}.docx")

                documento.save(caminho_saida)
                print(f"  ✔ Documento para '{nome_identificador}' gerado com sucesso!")
                total_gerado += 1

            except FileNotFoundError:
                print(f"\n[ERRO] Arquivo modelo '{caminho_modelo}' não foi encontrado!")
                break
            except KeyError as e:
                print(f"\n[ERRO] A coluna {e} não foi encontrada no Excel para o item {nome_identificador}.")
            except Exception as e:
                print(f"\n[ERRO] Ocorreu um erro ao gerar o documento para {nome_identificador}: {e}")
        else:
            print(f"\n[AVISO] O índice {index} é inválido e será ignorado.")

    print("\n--- Processo Finalizado! ---")
    print(f"{total_gerado} documento(s) foram gerados na pasta '{PASTA_SAIDA}'.")
    input("\nPressione Enter para sair.")


if __name__ == "__main__":
    motor_gerador()